import requests
import re
import fitz
import io
import logging
import concurrent.futures
import pytesseract
import trafilatura
from bs4 import BeautifulSoup
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from urllib.parse import urlparse
from datetime import datetime
from PIL import Image
from typing import List, Dict, Tuple, Optional
import os
import gc

# Настройка логгера
logging.basicConfig(
    filename='knowledge_base.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger()

# Конфигурация обработчиков
HEADERS = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
TIMEOUT = 40
MAX_WORKERS = 7
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200
OUTPUT_DIR = "processed_documents"  # Директория для сохранения документов

def ensure_output_directory():
    """Создает директорию для сохранения документов, если она не существует"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        logger.info(f"Создана директория для сохранения документов: {OUTPUT_DIR}")

def save_as_markdown(text: str, metadata: Dict, chunk_id: Optional[str] = None) -> str:
    """
    Сохраняет текст в формате Markdown с метаданными.
    Возвращает путь к сохраненному файлу.
    """
    ensure_output_directory()
    
    # Создаем имя файла из заголовка или URL
    title = metadata.get('title', 'document')
    source = urlparse(metadata['source']).path.replace('/', '_')[:100]
    clean_title = re.sub(r'[^\w\-_\. ]', '_', title)[:50]
    
    # Формируем имя файла
    if chunk_id:
        filename = f"{clean_title}_{chunk_id}.md"
    else:
        filename = f"{clean_title}_{source}.md"
    
    filepath = os.path.join(OUTPUT_DIR, filename)
    
    # Формируем Markdown контент с метаданными
    markdown_content = f"""# {title}\n\n"""
    markdown_content += f"**Источник:** [{metadata['source']}]({metadata['source']})\n\n"
    markdown_content += f"**Тип документа:** {metadata['type']}\n\n"
    markdown_content += f"**Дата обработки:** {metadata['processed_at']}\n\n"
    
    if 'author' in metadata:
        markdown_content += f"**Автор:** {metadata['author']}\n\n"
    if 'created' in metadata:
        markdown_content += f"**Дата создания:** {metadata['created']}\n\n"
    if 'modified' in metadata:
        markdown_content += f"**Дата изменения:** {metadata['modified']}\n\n"
    if 'pages' in metadata:
        markdown_content += f"**Количество страниц:** {metadata['pages']}\n\n"
    
    markdown_content += "---\n\n"
    markdown_content += text
    
    # Сохраняем файл
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        logger.info(f"Документ сохранен как Markdown: {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Ошибка сохранения Markdown файла {filepath}: {str(e)}")
        return ""

def download_document(url: str) -> Dict:
    """Загрузка документа с обработкой ошибок и таймаутом"""
    try:
        response = requests.get(url, headers=HEADERS, timeout=TIMEOUT)
        response.raise_for_status()
        
        content_type = response.headers.get('Content-Type', '')
        parsed_url = urlparse(url)
        extension = parsed_url.path.split('.')[-1].lower() if '.' in parsed_url.path else ''
        
        # Определение типа контента
        doc_type = 'unknown'
        if 'text/html' in content_type or extension in ('html', 'htm'):
            doc_type = 'html'
        elif 'application/pdf' in content_type or extension == 'pdf':
            doc_type = 'pdf'
        elif 'vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type or extension == 'docx':
            doc_type = 'docx'
        elif 'text/plain' in content_type or extension in ('txt', 'md'):
            doc_type = 'text'
        
        return {
            'content': response.content,
            'type': doc_type,
            'url': url,
            'status': 'success',
            'last_modified': response.headers.get('Last-Modified', '')
        }
    except requests.exceptions.RequestException as e:
        logger.error(f"Ошибка загрузки {url}: {str(e)}")
        return {'status': 'error', 'message': str(e), 'url': url}

def parse_html(content: bytes, url: str) -> Tuple[str, Dict]:
    """Парсинг HTML с улучшенным извлечением контента и поддержкой Markdown"""
    try:
        # Используем trafilatura для качественного извлечения текста в Markdown
        html_content = content.decode('utf-8', errors='replace')
        text = trafilatura.extract(
            html_content, 
            include_links=True,  # Сохраняем ссылки в Markdown формате
            include_tables=True,
            output_format='markdown'  # Получаем результат сразу в Markdown
        )
        
        if not text:
            # Fallback: BeautifulSoup если trafilatura не сработал
            soup = BeautifulSoup(html_content, 'html.parser')
            for element in soup(['script', 'style', 'header', 'footer', 'aside', 'nav']):
                element.decompose()
            
            # Конвертируем в Markdown-подобный формат
            text = ""
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table']):
                if element.name in ['h1', 'h2', 'h3']:
                    level = int(element.name[1])
                    text += f"{'#' * level} {element.get_text()}\n\n"
                elif element.name == 'p':
                    text += f"{element.get_text()}\n\n"
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        prefix = "- " if element.name == 'ul' else "1. "
                        text += f"{prefix}{li.get_text()}\n"
                    text += "\n"
                elif element.name == 'table':
                    text += "| " + " | ".join(th.get_text() for th in element.find_all('th')) + " |\n"
                    text += "| " + " | ".join(["---"] * len(element.find_all('th'))) + " |\n"
                    for row in element.find_all('tr'):
                        text += "| " + " | ".join(td.get_text() for td in row.find_all('td')) + " |\n"
                    text += "\n"
        
        # Извлечение метаданных
        soup = BeautifulSoup(html_content, 'html.parser') if 'soup' not in locals() else soup
        title = soup.title.string if soup.title else urlparse(url).path
        
        # Сохраняем оригинальный документ
        metadata = {
            'source': url,
            'type': 'html',
            'title': title,
            'processed_at': datetime.now().isoformat()
        }
        save_as_markdown(text, metadata)
        
        return text, metadata
    except Exception as e:
        logger.error(f"Ошибка парсинга HTML {url}: {str(e)}")
        return "", {}

def parse_pdf(content: bytes, url: str) -> Tuple[str, Dict]:
    """Парсинг PDF с обработкой OCR для сканированных документов и сохранением в Markdown"""
    try:
        text = ""
        metadata = {
            'source': url,
            'type': 'pdf',
            'processed_at': datetime.now().isoformat()
        }
        
        with fitz.open(stream=content, filetype="pdf") as pdf:
            metadata['pages'] = len(pdf)
            metadata['title'] = pdf.metadata.get('title', '') or os.path.basename(urlparse(url).path)
            
            for page_num, page in enumerate(pdf):
                page_text = page.get_text()
                
                # Если текст не извлекается, пробуем OCR
                if not page_text.strip():
                    try:
                        pix = page.get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        page_text = pytesseract.image_to_string(img, lang='rus+eng')
                        logger.info(f"Применен OCR для страницы {page_num+1} в {url}")
                    except Exception as ocr_error:
                        logger.warning(f"Ошибка OCR для {url} страница {page_num+1}: {str(ocr_error)}")
                
                # Форматируем текст страницы в Markdown
                text += f"\n\n## Страница {page_num+1}\n\n{page_text}"
        
        # Сохраняем оригинальный документ
        save_as_markdown(text, metadata)
        
        return text, metadata
    except Exception as e:
        logger.error(f"Ошибка парсинга PDF {url}: {str(e)}")
        return "", {}

def parse_docx(content: bytes, url: str) -> Tuple[str, Dict]:
    """Парсинг DOCX документов с конвертацией в Markdown"""
    try:
        docx = Document(io.BytesIO(content))
        text = ""
        
        # Обработка параграфов с сохранением структуры
        for para in docx.paragraphs:
            if not para.text.strip():
                continue
                
            # Определяем стиль для Markdown
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[1])
                text += f"{'#' * level} {para.text}\n\n"
            else:
                text += f"{para.text}\n\n"
        
        # Обработка таблиц в Markdown формате
        for table in docx.tables:
            text += "\n\n**Таблица:**\n\n"
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cells:
                    continue
                    
                if i == 0:  # Заголовок таблицы
                    text += "| " + " | ".join(cells) + " |\n"
                    text += "| " + " | ".join(["---"] * len(cells)) + " |\n"
                else:  # Данные таблицы
                    text += "| " + " | ".join(cells) + " |\n"
            text += "\n"
        
        # Извлечение метаданных
        metadata = {
            'source': url,
            'type': 'docx',
            'title': docx.core_properties.title or os.path.basename(urlparse(url).path),
            'author': docx.core_properties.author,
            'created': str(docx.core_properties.created),
            'modified': str(docx.core_properties.modified),
            'processed_at': datetime.now().isoformat()
        }
        
        # Сохраняем оригинальный документ
        save_as_markdown(text, metadata)
        
        return text, metadata
    except Exception as e:
        logger.error(f"Ошибка парсинга DOCX {url}: {str(e)}")
        return "", {}

def parse_text(content: bytes, url: str) -> Tuple[str, Dict]:
    """Обработка текстовых файлов с конвертацией в Markdown"""
    try:
        text = content.decode('utf-8', errors='replace')
        
        # Простая конвертация в Markdown (добавляем заголовок)
        markdown_text = f"# {os.path.basename(urlparse(url).path)}\n\n{text}"
        
        metadata = {
            'source': url,
            'type': 'text',
            'title': os.path.basename(urlparse(url).path),
            'processed_at': datetime.now().isoformat()
        }
        
        # Сохраняем оригинальный документ
        save_as_markdown(markdown_text, metadata)
        
        return markdown_text, metadata
    except Exception as e:
        logger.error(f"Ошибка парсинга текста {url}: {str(e)}")
        return "", {}

def clean_and_normalize(text: str) -> str:
    """Улучшенная очистка и нормализация текста с сохранением Markdown разметки"""
    # Сохраняем Markdown разметку
    markdown_elements = re.findall(r'(#+\s.*|\*\*.*\*\*|_.*_|~~.*~~|`.*`|\[.*\]\(.*\)|\!\[.*\]\(.*\)|^\s*[-*+]\s.*|^\s*\d+\.\s.*|^\s*>\s.*)', text, flags=re.MULTILINE)
    
    # Удаление лишних пробелов и переносов (кроме Markdown разметки)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.strip() in markdown_elements or any(markdown in line for markdown in markdown_elements):
            cleaned_lines.append(line)
        else:
            cleaned_line = re.sub(r'\s+', ' ', line).strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
    
    text = '\n'.join(cleaned_lines)
    
    # Удаление технических символов (кроме Markdown)
    text = re.sub(r'(?<!\\)[\x00-\x1F\x7F-\x9F]', '', text)
    
    # Удаление email (кроме Markdown ссылок)
    text = re.sub(r'(?<!\[)\S*@\S*\s?(?!.*\))', '', text)
    
    # Удаление URL (кроме Markdown ссылок)
    text = re.sub(r'(?<!\[)http\S+(?!.*\))', '', text)
    
    # Нормализация пунктуации (кроме Markdown)
    text = re.sub(r'([.!?])\s*(?![^[]*\))', r'\1 ', text)
    
    # Удаление повторяющихся пунктуаций (кроме Markdown)
    text = re.sub(r'([.!?])\1+(?![^[]*\))', r'\1', text)
    
    # Удаление специальных конструкций (кроме Markdown)
    text = re.sub(r'(\d+\.\d+\.\d+|\d+/\d+/\d+)(?![^[]*\))', '', text)  # Даты
    text = re.sub(r'(?<!\\)\[.*?(?<!\\)\]|(?<!\\)\{.*?(?<!\\)\}|(?<!\\)\(.*?(?<!\\)\)(?!\))', '', text)  # Квадратные/фигурные скобки
    
    # Удаление номеров страниц
    text = re.sub(r'\bСтраница\s+\d+\b', '', text, flags=re.IGNORECASE)
    
    return text

def process_single_document(url: str) -> List[Dict]:
    """Обработка одного документа с улучшенной обработкой ошибок"""
    try:
        logger.info(f"Начата обработка: {url}")
        doc = download_document(url)
        if doc['status'] != 'success':
            return []
        
        content = doc['content']
        doc_type = doc['type']
        
        # Распределение по типам документов
        if doc_type == 'html':
            text, metadata = parse_html(content, url)
        elif doc_type == 'pdf':
            text, metadata = parse_pdf(content, url)
        elif doc_type == 'docx':
            text, metadata = parse_docx(content, url)
        elif doc_type == 'text':
            text, metadata = parse_text(content, url)
        else:
            logger.warning(f"Неизвестный тип документа: {url}")
            return []
        
        if not text.strip():
            logger.warning(f"Пустой контент: {url}")
            return []
        
        cleaned_content = clean_and_normalize(text)
        
        # Разбиение на чанки
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        chunks = splitter.split_text(cleaned_content)
        
        # Сохранение чанков как отдельных Markdown файлов
        chunk_files = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_id'] = f"{url}-{i}"
            chunk_metadata['chunk_length'] = len(chunk)
            
            # Сохраняем чанк
            chunk_file = save_as_markdown(
                text=chunk,
                metadata=chunk_metadata,
                chunk_id=f"chunk_{i}"
            )
            if chunk_file:
                chunk_files.append(chunk_file)
        
        # Логирование информации о сохраненных файлах
        if chunk_files:
            logger.info(f"Сохранены чанки документа {url} в файлы: {', '.join(chunk_files)}")
        
        # Формирование результатов
        result = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_id'] = f"{url}-{i}"
            chunk_metadata['chunk_length'] = len(chunk)
            result.append({'text': chunk, 'metadata': chunk_metadata})
        
        logger.info(f"Успешно обработан: {url} - {len(chunks)} чанков")
        return result
    
    except Exception as e:
        logger.error(f"Критическая ошибка обработки {url}: {str(e)}", exc_info=True)
        return []
    finally:
        # Освобождение памяти
        if 'content' in locals():
            del content
        gc.collect()

def generate_processing_report(all_chunks: List[Dict], urls: List[str], report_file: str = "processing_report.md") -> str:
    """
    Генерирует подробный отчет о качестве и количестве обработанных документов в формате Markdown.
    Возвращает путь к сохраненному файлу отчета.
    """
    ensure_output_directory()
    
    # Собираем статистику
    total_urls = len(urls)
    processed_sources = set(chunk['metadata']['source'] for chunk in all_chunks)
    total_processed = len(processed_sources)
    total_chunks = len(all_chunks)
    avg_chunk_size = sum(len(chunk['text']) for chunk in all_chunks) // total_chunks if total_chunks > 0 else 0
    
    # Анализ качества
    empty_chunks = sum(1 for chunk in all_chunks if not chunk['text'].strip())
    small_chunks = sum(1 for chunk in all_chunks if 0 < len(chunk['text'].strip()) < 50)
    
    # Анализ по типам документов
    doc_types = {}
    for chunk in all_chunks:
        doc_type = chunk['metadata']['type']
        doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
    
    # Формируем Markdown отчет
    report = "# Отчет о обработке документов\n\n"
    report += f"**Дата генерации отчета:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    
    report += "## Общая статистика\n\n"
    report += f"- Всего URL для обработки: {total_urls}\n"
    report += f"- Успешно обработано документов: {total_processed} ({total_processed/total_urls:.1%})\n"
    report += f"- Не удалось обработать: {total_urls - total_processed}\n"
    report += f"- Всего чанков: {total_chunks}\n"
    report += f"- Средний размер чанка: {avg_chunk_size} символов\n\n"
    
    report += "## Анализ качества чанков\n\n"
    report += f"- Пустых чанков: {empty_chunks} ({empty_chunks/total_chunks:.1%})\n"
    report += f"- Маленьких чанков (<50 символов): {small_chunks} ({small_chunks/total_chunks:.1%})\n\n"
    
    report += "## Распределение по типам документов\n\n"
    for doc_type, count in doc_types.items():
        report += f"- {doc_type.capitalize()}: {count} чанков ({count/total_chunks:.1%})\n"
    report += "\n"
    
    report += "## Список обработанных источников\n\n"
    for i, source in enumerate(sorted(processed_sources), 1):
        report += f"{i}. [{source}]({source})\n"
    
    # Сохраняем отчет
    full_report_path = os.path.join(OUTPUT_DIR, report_file)
    with open(full_report_path, 'w', encoding='utf-8') as f:
        f.write(report)
    
    logger.info(f"Отчет о обработке сохранен в {full_report_path}")
    return full_report_path

def create_knowledge_base(
    urls: List[str],
    db_path: str = "medical_knowledge_db",
    embedding_model: str = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
) -> FAISS:
    """Создание базы знаний из списка URL"""
    logger.info(f"Начато создание базы знаний из {len(urls)} документов")
    all_chunks = []
    
    # Многопоточная обработка
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_url = {executor.submit(process_single_document, url): url for url in urls}
        for future in concurrent.futures.as_completed(future_to_url):
            url = future_to_url[future]
            try:
                chunks = future.result()
                all_chunks.extend(chunks)
            except Exception as e:
                logger.error(f"Ошибка при обработке {url}: {str(e)}")
    
    if not all_chunks:
        logger.error("Нет данных для создания базы знаний")
        raise ValueError("Не удалось обработать ни один документ")
    
    # Подготовка данных для векторизации
    texts = [chunk['text'] for chunk in all_chunks]
    metadatas = [chunk['metadata'] for chunk in all_chunks]
    
    logger.info(f"Создание эмбеддингов с моделью {embedding_model}")
    embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
    
    logger.info("Построение векторной базы")
    vector_db = FAISS.from_texts(
        texts=texts,
        embedding=embeddings,
        metadatas=metadatas
    )
    
    # Сохранение базы знаний
    vector_db.save_local(db_path)
    logger.info(f"База знаний создана: {len(all_chunks)} чанков, сохранено в {db_path}")
    
    # Генерация отчета
    report = f"Отчет по созданию базы знаний:\n"
    report += f"- Обработано URL: {len(urls)}\n"
    report += f"- Успешно обработано документов: {len(set(chunk['metadata']['source'] for chunk in all_chunks))}\n"
    report += f"- Всего чанков: {len(all_chunks)}\n"
    report += f"- Средний размер чанка: {sum(len(t) for t in texts)//len(texts)} символов\n"
    
    # Сохраняем отчет в Markdown
    report_file = os.path.join(OUTPUT_DIR, "processing_report.md")
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write(report)
    
    logger.info(report)
    logger.info(f"Отчет сохранен в {report_file}")
        # Генерация и сохранение отчета о качестве обработки
    generate_processing_report(all_chunks, urls)
    
    return vector_db


if __name__ == "__main__":
    # Список медицинских документов для обработки
    medical_docs = [
        "https://minzdrav.gov.ru/documents/7025-federalnyy-zakon-323-fz-ot-21-noyabrya-2011-g",
        "https://www.garant.ru/products/ipo/prime/doc/405742919/",
        "https://base.garant.ru/71431038/",
        "https://docs.cntd.ru/document/542615520",
        "https://base.garant.ru/72147566/#:~:text=декабря%202018%20г.-,N%20898н%20%22О%20внесении%20изменений%20в%20сроки%20и%20этапы%20аккредитации,от%2022%20декабря%202017%20г.",
        "https://minzdrav.gov.ru/documents/8956-prikaz-ministerstva-zdravoohraneniya-rossiyskoy-federatsii-ot-3-avgusta-2012-g-66n-ob-utverzhdenii-poryadka-i-srokov-sovershenstvovaniya-meditsinskimi-rabotnikami-i-farmatsevticheskimi-rabotnikami-professionalnyh-znaniy-i-navykov-putem-obucheniya-po-dopolnitelnym-professionalnym-obrazovatelnym-programmam-v-obrazovatelnyh-i-nauchnyh-organizatsiyah",
        "https://docs.cntd.ru/document/420310213",
        "https://normativ.kontur.ru/document?moduleId=1&documentId=269836",
        "https://mosgorzdrav.ru/ru-RU/moscowDoctorDetail.html",
        "https://normativ.kontur.ru/document?moduleId=1&documentId=457082",
        "http://publication.pravo.gov.ru/Document/View/0001202101120007?index=2"
    ]
    
    # Создание базы знаний
    knowledge_base = create_knowledge_base(
        urls=medical_docs,
        db_path="medical_regulations_db",
        embedding_model="sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
    )
    
    # Пример использования базы знаний
    query = "требования к аккредитации медицинских работников"
    similar_docs = knowledge_base.similarity_search(query, k=3)
    
    print(f"\nРезультаты поиска по запросу: '{query}'")
    for i, doc in enumerate(similar_docs):
        print(f"\nДокумент {i+1}:")
        print(f"Источник: {doc.metadata['source']}")
        print(f"Тип: {doc.metadata['type']}")
        print(f"Заголовок: {doc.metadata.get('title', 'без названия')}")
        print(f"Содержание:\n{doc.page_content[:300]}...")
        